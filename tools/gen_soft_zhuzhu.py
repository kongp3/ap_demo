import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

SOFTWARE_NAME = '智能审计系统'
VERSION = 'V1.0'
MAX_LINES = 3200

# 支持的源码文件后缀
CODE_SUFFIX = ['.py', '.js', '.ts', '.vue', '.html', '.css']

# 注释正则
COMMENT_PATTERNS = {
    '.py': re.compile(r'^\s*#'),
    '.js': re.compile(r'^\s*//'),
    '.ts': re.compile(r'^\s*//'),
    '.vue': re.compile(r'^\s*//|^\s*<!--|^\s*\*|^\s*<\!--'),
    '.html': re.compile(r'^\s*<!--'),
    '.css': re.compile(r'^\s*/\*|^\s*\*'),
}

# 敏感信息正则
SENSITIVE_PATTERNS = [
    re.compile(r'(?i)api[_-]?key\s*[:=]\s*[\"\"][^\"\"]+[\"\"]'),
    re.compile(r'1[3-9]\d{9}'), # 手机号
]

EXCLUDE_DIRS = {'node_modules', 'public', '.vscode', '__pycache__'}

def is_comment(line, ext):
    pat = COMMENT_PATTERNS.get(ext, None)
    return bool(pat and pat.match(line))

def is_blank(line):
    return not line.strip()

def remove_sensitive(line):
    for pat in SENSITIVE_PATTERNS:
        line = pat.sub('"***"', line)
    return line

def collect_code_files(root_dirs):
    files = []
    for root_dir in root_dirs:
        for dirpath, dirnames, filenames in os.walk(root_dir):
            dirnames[:] = [d for d in dirnames if d not in EXCLUDE_DIRS]
            for f in filenames:
                ext = os.path.splitext(f)[1]
                if ext in CODE_SUFFIX:
                    files.append(os.path.join(dirpath, f))
    return sorted(files)

def extract_valid_lines(filepath):
    ext = os.path.splitext(filepath)[1]
    valid_lines = []
    with open(filepath, encoding='utf-8', errors='ignore') as f:
        for line in f:
            if is_blank(line) or is_comment(line, ext):
                continue
            line = remove_sensitive(line.rstrip('\n'))
            valid_lines.append(line)
    return valid_lines

START_LINE_NUM = 1  # 新增：自定义起始行号

def main():
    # 只遍历后端app和前端src目录
    code_files = collect_code_files(['app', 'app/static/src'])
    all_lines = []
    for file in code_files:
        valid = extract_valid_lines(file)
        if valid:
            all_lines.append(f'// 文件: {file}')
            all_lines.extend(valid)
    # 过滤掉空内容的代码行（不含文件标记）
    all_lines = [line for line in all_lines if line.strip() and not (re.match(r'^\d+:\s*$', line))]
    # 只保留MAX_LINES行有效代码（不含文件标记）
    code_lines = [line for line in all_lines if not line.startswith('// 文件:')]
    if len(code_lines) > MAX_LINES:
        code_lines = code_lines[:MAX_LINES]
    # 写入docx
    doc = Document()
    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(9)  # 小五号
    # 页眉
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = f'{SOFTWARE_NAME}{VERSION} 源代码节选'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 代码行，带连续行号，起始为START_LINE_NUM
    for idx, line in enumerate(code_lines, START_LINE_NUM):
        code_p = doc.add_paragraph()
        code_p.paragraph_format.left_indent = Pt(12)
        code_p.style = style
        code_p.add_run(f'{idx}: {line}')
    doc.save('soft_zhuzhu.docx')
    print(f'已生成 soft_zhuzhu.docx, 共{len(code_lines)}行，起始行号为{START_LINE_NUM}')

if __name__ == '__main__':
    main()
